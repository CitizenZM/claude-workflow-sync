#!/usr/bin/env python3
"""Generate Chinese Affiliate Weekly Report (.docx) for one region (US or EU).

Usage:
    python3 generate_report_zh.py --region us --out /path/Awin-Rockbros-US-周报.docx
    python3 generate_report_zh.py --region eu --out /path/Awin-Rockbros-EU-周报.docx

Embeds matplotlib charts (trend line, channel donut, top-publisher bar,
concentration gauge) plus screenshots from output/. All section titles,
prose, table headers and diagnoses/suggestions are in Chinese.
"""
import argparse
import json
import os
import sys
from datetime import date
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from matplotlib import font_manager

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --------------------------- Style constants ---------------------------
NAVY = RGBColor(0x0B, 0x2E, 0x4F)
ACCENT = RGBColor(0xE7, 0x4C, 0x3C)
GOLD = RGBColor(0xC9, 0x8A, 0x2B)
GREEN = RGBColor(0x2C, 0x8A, 0x4F)
GREY = RGBColor(0x55, 0x5F, 0x6D)
LIGHT_HEX = "F4F6F8"
HEADER_HEX = "0B2E4F"
ACCENT_HEX = "E74C3C"
DIAG_HEX = "FFF8E1"
ZH_FONT = "PingFang SC"   # docx Chinese face — Word on macOS resolves it
ZH_FONT_FALLBACK = "Microsoft YaHei"
EN_FONT = "Calibri"

# matplotlib Chinese font setup
for cand in ["PingFang SC", "Heiti TC", "Arial Unicode MS", "Songti SC"]:
    if any(f.name == cand for f in font_manager.fontManager.ttflist):
        plt.rcParams["font.family"] = cand
        break
plt.rcParams["axes.unicode_minus"] = False
plt.rcParams["axes.spines.top"] = False
plt.rcParams["axes.spines.right"] = False

ARTIFACT_DIR = Path("/Users/xiaozuo/.claude/skills/awin-rockbros-weekly-report/output")
CHART_DIR = ARTIFACT_DIR / "charts"
CHART_DIR.mkdir(exist_ok=True)


# --------------------------- docx helpers ---------------------------
def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_run_zh(run, text, size=11, bold=False, italic=False, color=None):
    run.text = text
    run.bold = bold
    run.italic = italic
    run.font.name = EN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), ZH_FONT)
    rfonts.set(qn("w:ascii"), EN_FONT)
    rfonts.set(qn("w:hAnsi"), EN_FONT)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, bold=False, size=11, color=None, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run("")
    set_run_zh(r, text, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1, color=NAVY):
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("")
    set_run_zh(r, text, size=sizes.get(level, 12), bold=True, color=color)
    # Add bottom border for level 1
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), HEADER_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_bullet(doc, text, size=11, color=None):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("")
    set_run_zh(r, text, size=size, color=color)
    return p


def add_table(doc, headers, rows, col_widths=None, accent_first_col=False):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run("")
        set_run_zh(run, h, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, HEADER_HEX)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run("")
            is_first = (ci == 0 and accent_first_col)
            set_run_zh(run, str(val), size=10,
                       bold=is_first,
                       color=NAVY if is_first else None)
            if ri % 2 == 0:
                shade_cell(cell, LIGHT_HEX)
    if col_widths:
        for ri in range(len(table.rows)):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = w
    return table


def add_diagnosis(doc, diagnosis, suggestions):
    """Yellow callout box with 诊断 + 建议."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, DIAG_HEX)
    cell.text = ""
    # Title
    p1 = cell.paragraphs[0]
    r = p1.add_run("")
    set_run_zh(r, "🔍 诊断与建议", size=11, bold=True, color=GOLD)
    # Diagnosis
    p2 = cell.add_paragraph()
    r = p2.add_run("")
    set_run_zh(r, "诊断：", size=10, bold=True, color=NAVY)
    r2 = p2.add_run("")
    set_run_zh(r2, diagnosis, size=10, color=GREY)
    # Suggestions
    p3 = cell.add_paragraph()
    r = p3.add_run("")
    set_run_zh(r, "建议：", size=10, bold=True, color=NAVY)
    for i, sug in enumerate(suggestions, 1):
        psug = cell.add_paragraph()
        psug.paragraph_format.left_indent = Cm(0.5)
        r = psug.add_run("")
        set_run_zh(r, f"  {i}. {sug}", size=10, color=GREY)
    doc.add_paragraph()


def add_screenshot(doc, path, caption=None, width_in=6.2):
    if not path or not os.path.exists(path):
        add_para(doc, f"[截图缺失：{path}]", italic=True, color=GREY, size=9)
        return
    try:
        doc.add_picture(path, width=Inches(width_in))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run("")
            set_run_zh(r, caption, size=9, italic=True, color=GREY)
    except Exception as exc:
        add_para(doc, f"[截图嵌入失败：{exc}]", italic=True, color=GREY, size=9)


# --------------------------- Chart generators ---------------------------
def chart_gmv_trend(region, days, prev, this_week, currency, out_path):
    fig, ax = plt.subplots(figsize=(8.4, 3.4))
    x = list(range(len(days)))
    ax.plot(x, prev, marker="o", linewidth=2, color="#9CA3AF", label="上周")
    ax.plot(x, this_week, marker="o", linewidth=2.6, color="#0B2E4F", label="本周")
    ax.fill_between(x, this_week, alpha=0.08, color="#0B2E4F")
    ax.set_xticks(x)
    ax.set_xticklabels(days, fontsize=9)
    ax.set_title(f"{region} — 七日 GMV 趋势对比 ({currency})", fontsize=12, color="#0B2E4F", weight="bold", pad=10)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{currency}{v:,.0f}"))
    ax.grid(True, axis="y", linestyle="--", alpha=0.35)
    ax.legend(loc="upper right", frameon=False, fontsize=9)
    for i, v in enumerate(this_week):
        ax.annotate(f"{currency}{v:,.0f}", (i, v), textcoords="offset points",
                    xytext=(0, 8), ha="center", fontsize=8, color="#0B2E4F")
    plt.tight_layout()
    plt.savefig(out_path, dpi=160, bbox_inches="tight")
    plt.close()


def chart_channel_donut(region, channels, values, out_path):
    fig, ax = plt.subplots(figsize=(5.0, 4.2))
    colors = ["#0B2E4F", "#E74C3C", "#C98A2B", "#2C8A4F", "#7C3AED", "#0EA5E9"]
    wedges, texts, autotexts = ax.pie(
        values, labels=channels, autopct="%1.1f%%",
        startangle=90, pctdistance=0.78,
        colors=colors[:len(values)],
        wedgeprops=dict(width=0.42, edgecolor="white", linewidth=2),
        textprops=dict(fontsize=10, color="#0B2E4F"),
    )
    for t in autotexts:
        t.set_color("white"); t.set_fontsize(9); t.set_weight("bold")
    ax.set_title(f"{region} — 渠道结构 (Channel Mix)", fontsize=12, color="#0B2E4F", weight="bold", pad=12)
    plt.tight_layout()
    plt.savefig(out_path, dpi=160, bbox_inches="tight")
    plt.close()


def chart_top_publishers(region, names, values, currency, out_path):
    fig, ax = plt.subplots(figsize=(8.4, max(3.0, 0.32 * len(names) + 1.2)))
    y = list(range(len(names)))
    bars = ax.barh(y, values, color="#0B2E4F", edgecolor="white")
    ax.set_yticks(y)
    ax.set_yticklabels(names, fontsize=9)
    ax.invert_yaxis()
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f"{currency}{v:,.0f}"))
    ax.set_title(f"{region} — Top 10 Publisher GMV", fontsize=12, color="#0B2E4F", weight="bold", pad=10)
    ax.grid(True, axis="x", linestyle="--", alpha=0.35)
    for bar, v in zip(bars, values):
        ax.text(v, bar.get_y() + bar.get_height()/2,
                f"  {currency}{v:,.0f}", va="center", fontsize=9, color="#0B2E4F")
    plt.tight_layout()
    plt.savefig(out_path, dpi=160, bbox_inches="tight")
    plt.close()


def chart_funnel(region, stages, counts, out_path):
    fig, ax = plt.subplots(figsize=(7.5, 3.4))
    colors = ["#0B2E4F", "#1E5A99", "#3B82F6", "#60A5FA", "#93C5FD"]
    bars = ax.barh(range(len(stages)), counts, color=colors[:len(stages)], edgecolor="white")
    ax.set_yticks(range(len(stages)))
    ax.set_yticklabels(stages, fontsize=10)
    ax.invert_yaxis()
    ax.set_title(f"{region} — Publisher 生命周期漏斗", fontsize=12, color="#0B2E4F", weight="bold", pad=10)
    ax.grid(True, axis="x", linestyle="--", alpha=0.35)
    for bar, c in zip(bars, counts):
        ax.text(c, bar.get_y() + bar.get_height()/2,
                f"  {c:,}", va="center", fontsize=10, color="#0B2E4F", weight="bold")
    plt.tight_layout()
    plt.savefig(out_path, dpi=160, bbox_inches="tight")
    plt.close()


def chart_concentration_gauge(region, top3_pct, out_path):
    fig, ax = plt.subplots(figsize=(5.0, 3.0), subplot_kw={"projection": "polar"})
    # Half donut from -90 → 90 deg
    import numpy as np
    theta = np.linspace(np.pi, 0, 100)
    radii = [1] * 100
    # Three zones: green (<40%), amber (40-60%), red (>60%)
    zones = [(np.pi, np.pi*0.6, "#2C8A4F"),
             (np.pi*0.6, np.pi*0.4, "#C98A2B"),
             (np.pi*0.4, 0, "#E74C3C")]
    for start, end, c in zones:
        t = np.linspace(start, end, 50)
        ax.fill_between(t, 0.85, 1.0, color=c, alpha=0.85)
    needle_angle = np.pi * (1 - top3_pct / 100)
    ax.plot([needle_angle, needle_angle], [0, 0.95], color="#0B2E4F", linewidth=4)
    ax.scatter([needle_angle], [0.95], color="#0B2E4F", s=80, zorder=5)
    ax.set_ylim(0, 1.1)
    ax.set_xticks([np.pi, np.pi*0.75, np.pi*0.5, np.pi*0.25, 0])
    ax.set_xticklabels(["0%", "25%", "50%", "75%", "100%"], fontsize=9)
    ax.set_yticks([])
    ax.set_title(f"{region} — Top3 集中度  {top3_pct:.1f}%",
                 fontsize=12, color="#0B2E4F", weight="bold", pad=18)
    ax.grid(False)
    ax.spines["polar"].set_visible(False)
    plt.tight_layout()
    plt.savefig(out_path, dpi=160, bbox_inches="tight")
    plt.close()


# --------------------------- Cover ---------------------------
def build_cover(doc, data):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("")
    set_run_zh(r, "ROCKBROS GLOBAL", size=14, bold=True, color=ACCENT)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("")
    set_run_zh(r, f"{data['region_name']}联盟营销周报", size=28, bold=True, color=NAVY)

    rng = doc.add_paragraph()
    rng.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rng.add_run("")
    set_run_zh(r, f"报告周期截止 {data['report_date']}  •  {data['region_name']} (Awin {data['merchant_id']})",
               size=12, color=GREY)

    doc.add_paragraph()
    add_para(doc, "汇报对象：", bold=True, size=10, color=GREY)
    add_para(doc, "CEO / CMO / 增长负责人  ──  阅读时间 ≤ 2 分钟（执行摘要部分）", size=10, color=GREY)
    add_para(doc, "汇报方：", bold=True, size=10, color=GREY)
    add_para(doc, "Cell Digital Technology Inc. — 联盟营销绩效办公室", size=10, color=GREY)
    add_para(doc, f"数据来源：Awin 广告主后台（{data['region_name']} {data['merchant_id']}），数据捕获日期 {data['report_date']}。",
             size=10, color=GREY, italic=True)

    doc.add_paragraph()
    add_para(doc, "—  报告框架  —", bold=True, size=11, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    toc = [
        "1. 执行摘要 (Executive Summary)",
        "2. GMV 与绩效深度分析",
        "3. Publisher 表现分析",
        "4. 本周新增 Publisher（招募与激活）",
        "5. 活动 / 推广绩效",
        "6. 内容与达人表现",
        "7. 风险与异常",
        "8. 下周行动计划",
        "9. KPI 仪表盘",
    ]
    for line in toc:
        add_para(doc, f"   {line}", size=11, color=NAVY)
    doc.add_page_break()


# --------------------------- Sections ---------------------------
def section_executive(doc, data):
    add_heading(doc, "一、执行摘要 (Executive Summary)", level=1)
    add_para(doc, "面向 CEO/CMO/增长负责人 — 全周联盟营销绩效快照、环比变动、以及三条最关键判断。",
             italic=True, color=GREY, size=10)

    add_heading(doc, "1.1 关键 KPI 快照", level=2)
    add_table(doc, ["指标", "本周", "环比 WoW", "目标", "状态"], data["exec_kpi"], accent_first_col=True)

    doc.add_paragraph()
    add_heading(doc, "1.2 三大核心判断", level=2)
    for tk in data["exec_takeaways"]:
        add_bullet(doc, tk)

    doc.add_paragraph()
    add_diagnosis(doc, data["exec_diagnosis"], data["exec_suggestions"])
    doc.add_page_break()


def section_gmv(doc, data):
    add_heading(doc, "二、GMV 与绩效深度分析", level=1)

    add_heading(doc, "2.1 七日 GMV 趋势（本周 vs 上周）", level=2)
    add_screenshot(doc, str(CHART_DIR / f"{data['region_code']}_gmv_trend.png"),
                   caption=f"{data['region_name']} — 本周 vs 上周日 GMV 走势")

    add_table(doc, ["日期", "本周 GMV", "上周 GMV", "环比", "备注"], data["gmv_trend"])

    doc.add_paragraph()
    add_heading(doc, "2.2 Awin 后台原始截图", level=2)
    if data.get("home_screenshot"):
        add_screenshot(doc, data["home_screenshot"],
                       caption=f"Awin {data['region_name']} ({data['merchant_id']}) — 首页 KPI 看板")

    add_heading(doc, "2.3 渠道结构（Channel Mix）", level=2)
    add_screenshot(doc, str(CHART_DIR / f"{data['region_code']}_channel_donut.png"),
                   caption=f"{data['region_name']} — 渠道 GMV 占比")
    add_table(doc, ["渠道类型", "GMV", "占比", "环比", "备注"], data["channel_breakdown"], accent_first_col=True)
    add_para(doc, "重点：绩效型渠道（Cashback / Coupon）与品牌型渠道（Content / Creator）的平衡。"
                  "目标内容型 GMV 占比 30–50%。", italic=True, color=GREY, size=10)

    add_heading(doc, "2.4 地域 / 设备 / 品类拆解", level=2)
    add_table(doc, ["维度", "Top 驱动因子", "洞察"], data["geo_device_product"], accent_first_col=True)

    doc.add_paragraph()
    add_diagnosis(doc, data["gmv_diagnosis"], data["gmv_suggestions"])
    doc.add_page_break()


def section_publishers(doc, data):
    add_heading(doc, "三、Publisher 表现分析", level=1)

    add_heading(doc, "3.1 Top 10 Publisher GMV", level=2)
    add_screenshot(doc, str(CHART_DIR / f"{data['region_code']}_top_pub.png"),
                   caption=f"{data['region_name']} — Top 10 Publisher 周 GMV 排名")
    add_table(doc, ["Publisher", "区域", "GMV", "环比", "类型", "动作"], data["top_publishers"], accent_first_col=True)

    if data.get("publishers_screenshot"):
        doc.add_paragraph()
        add_screenshot(doc, data["publishers_screenshot"],
                       caption=f"Awin {data['region_name']} — Publisher 列表全景截图")

    doc.add_paragraph()
    add_heading(doc, "3.2 集中度风险 (Concentration Risk)", level=2)
    add_screenshot(doc, str(CHART_DIR / f"{data['region_code']}_gauge.png"),
                   caption=f"{data['region_name']} — Top 3 Publisher GMV 占比仪表盘")
    add_table(doc, ["指标", "数值", "基准", "风险等级"], data["concentration"], accent_first_col=True)

    add_heading(doc, "3.3 Publisher 生命周期漏斗", level=2)
    add_screenshot(doc, str(CHART_DIR / f"{data['region_code']}_funnel.png"),
                   caption=f"{data['region_name']} — Publisher 招募 → 激活 → 转化 漏斗")
    add_table(doc, ["阶段", "数量", "转化率"], data["lifecycle"], accent_first_col=True)
    add_para(doc, "诊断维度：低激活 = 入驻摩擦/沟通缺失；低转化 = 报价或内容匹配度不足。",
             italic=True, color=GREY, size=10)

    doc.add_paragraph()
    add_diagnosis(doc, data["pub_diagnosis"], data["pub_suggestions"])
    doc.add_page_break()


def section_new_publishers(doc, data):
    add_heading(doc, "四、本周新增 Publisher（招募与激活）", level=1)
    add_para(doc, f"过去 7 日新增合作 Publisher 共 {data['new_publishers_count']} 家，"
                  f"以下为详细名单（按加入时间倒序）。",
             size=11, color=NAVY)

    add_heading(doc, "4.1 新增 Publisher 列表", level=2)
    add_table(doc, ["Publisher", "网站", "推广类型", "行业", "加入日期"],
              data["new_publishers"], accent_first_col=True)

    add_heading(doc, "4.2 新增类型分布", level=2)
    add_table(doc, ["推广类型", "新增数", "占比", "下一步动作"],
              data["new_pub_breakdown"], accent_first_col=True)

    add_heading(doc, "4.3 激活跟踪 (TTFS — Time To First Sale)", level=2)
    add_table(doc, ["Publisher", "状态", "首条内容", "GMV", "下一步"],
              data["activation"], accent_first_col=True)

    doc.add_paragraph()
    add_diagnosis(doc, data["recruit_diagnosis"], data["recruit_suggestions"])
    doc.add_page_break()


def section_campaigns(doc, data):
    add_heading(doc, "五、活动 / 推广绩效", level=1)

    add_heading(doc, "5.1 活动结果", level=2)
    add_table(doc, ["活动", "类型", "GMV", "ROI", "备注"], data["campaigns"], accent_first_col=True)

    add_heading(doc, "5.2 报价/Offer 表现", level=2)
    add_table(doc, ["Offer 类型", "CVR", "AOV", "洞察"], data["offers"], accent_first_col=True)

    doc.add_paragraph()
    add_diagnosis(doc, data["camp_diagnosis"], data["camp_suggestions"])
    doc.add_page_break()


def section_content(doc, data):
    add_heading(doc, "六、内容与达人表现", level=1)

    add_heading(doc, "6.1 内容型 GMV 拆解", level=2)
    add_table(doc, ["内容类型", "GMV", "占比", "环比"], data["content_mix"], accent_first_col=True)

    add_heading(doc, "6.2 Top 内容产出", level=2)
    add_table(doc, ["创作者", "平台", "曝光", "GMV", "钩子/Hook"], data["top_content"], accent_first_col=True)

    doc.add_paragraph()
    add_diagnosis(doc, data["content_diagnosis"], data["content_suggestions"])
    doc.add_page_break()


def section_risks(doc, data):
    add_heading(doc, "七、风险与异常 (Issues & Risks)", level=1)
    add_table(doc, ["问题", "影响", "根因", "动作"], data["risks"], accent_first_col=True)

    doc.add_paragraph()
    add_diagnosis(doc, data["risk_diagnosis"], data["risk_suggestions"])
    doc.add_page_break()


def section_actions(doc, data):
    add_heading(doc, "八、下周行动计划 (Action Plan)", level=1)
    add_para(doc, "整份报告最重要的章节 — 每条动作均含责任人与可衡量结果。",
             italic=True, color=GREY, size=10)

    add_heading(doc, "8.1 增长动作 (Growth)", level=2)
    for a in data["actions_growth"]:
        add_bullet(doc, a)

    add_heading(doc, "8.2 招募计划 (Recruitment)", level=2)
    for a in data["actions_recruit"]:
        add_bullet(doc, a)

    add_heading(doc, "8.3 优化动作 (Optimization)", level=2)
    for a in data["actions_optimize"]:
        add_bullet(doc, a)
    doc.add_page_break()


def section_kpi(doc, data):
    add_heading(doc, "九、KPI 仪表盘 (Operator View)", level=1)
    add_table(doc, ["KPI", "当前", "目标", "差距", "动作"], data["kpi_dashboard"], accent_first_col=True)
    add_para(doc, "颜色规则：绿 = 达标/超额；黄 = 10% 以内差距；红 = 大于 10% 差距。",
             italic=True, color=GREY, size=10)
    add_para(doc, f"报告生成日期 {data['report_date']}  •  awin-rockbros-weekly-report skill  •  Cell Digital Technology Inc.",
             italic=True, color=GREY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)


# --------------------------- Build ---------------------------
def _parse_money(val, cur):
    s = str(val).replace(cur, "").replace(",", "").replace("€", "").replace("$", "").strip()
    if not s or s in ("—", "-", "N/A", "NA", "新"):
        return 0.0
    try:
        return float(s)
    except ValueError:
        return 0.0


def _parse_pct(val):
    s = str(val).replace("%", "").strip()
    try:
        return float(s)
    except ValueError:
        return 0.0


def render_charts(data):
    region = data["region_name"]
    code = data["region_code"]
    cur = data["currency_symbol"]

    days = [d[0] for d in data["gmv_trend"]]
    this_week = [_parse_money(d[1], cur) for d in data["gmv_trend"]]
    prev_week = [_parse_money(d[2], cur) for d in data["gmv_trend"]]
    chart_gmv_trend(region, days, prev_week, this_week, cur,
                    str(CHART_DIR / f"{code}_gmv_trend.png"))

    chans = [r[0] for r in data["channel_breakdown"]]
    pcts = [_parse_pct(r[2]) for r in data["channel_breakdown"]]
    chart_channel_donut(region, chans, pcts,
                        str(CHART_DIR / f"{code}_channel_donut.png"))

    # Top 10 publishers — filter to those with numeric GMV for the chart
    enriched = [(t[0], _parse_money(t[2], cur)) for t in data["top_publishers"][:10]]
    enriched_nonzero = [e for e in enriched if e[1] > 0]
    if not enriched_nonzero:
        enriched_nonzero = enriched
    names = [e[0] for e in enriched_nonzero]
    vals = [e[1] for e in enriched_nonzero]
    chart_top_publishers(region, names, vals, cur,
                         str(CHART_DIR / f"{code}_top_pub.png"))

    stages = [r[0] for r in data["lifecycle"]]
    counts = [int(str(r[1]).replace(",", "")) for r in data["lifecycle"]]
    chart_funnel(region, stages, counts, str(CHART_DIR / f"{code}_funnel.png"))

    chart_concentration_gauge(region, data["top3_concentration_pct"],
                              str(CHART_DIR / f"{code}_gauge.png"))


def build(data, out_path):
    render_charts(data)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
    build_cover(doc, data)
    section_executive(doc, data)
    section_gmv(doc, data)
    section_publishers(doc, data)
    section_new_publishers(doc, data)
    section_campaigns(doc, data)
    section_content(doc, data)
    section_risks(doc, data)
    section_actions(doc, data)
    section_kpi(doc, data)
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--data", required=True, help="path to region data JSON")
    ap.add_argument("--out", required=True, help="output .docx path")
    args = ap.parse_args()
    data = json.loads(Path(args.data).read_text())
    data.setdefault("report_date", date.today().isoformat())
    build(data, args.out)
    print(f"Wrote {args.out}")


if __name__ == "__main__":
    main()

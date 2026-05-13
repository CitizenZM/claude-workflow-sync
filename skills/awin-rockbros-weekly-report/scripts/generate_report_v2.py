#!/usr/bin/env python3
"""Awin Rockbros weekly report — v2 generator.

Changes from v1:
  - Drop sections 2.4 (Geo/Device/Product), 3.2 (Concentration), 3.3 (Lifecycle), 7 (Risks)
  - Skip any section whose table is empty (no placeholder rows)
  - Status indicators are real color dots (🟢🟡🔴) — already in data
  - GMV section pulls weekly + monthly + WoW + MoM
  - Channel mix uses real Awin "Publisher Performance by Quantity" %
  - New publishers limited to last 7 days only
"""
import argparse, json, datetime, sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Optional matplotlib (charts)
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import font_manager
    HAS_PLT = True
    # Pick a CJK-capable font
    for fam in ["PingFang SC", "Heiti TC", "Arial Unicode MS", "Songti SC", "STHeiti"]:
        try:
            font_manager.findfont(fam, fallback_to_default=False)
            plt.rcParams["font.family"] = fam
            break
        except Exception: pass
    plt.rcParams["axes.unicode_minus"] = False
except Exception:
    HAS_PLT = False

CHART_DIR = Path(__file__).resolve().parent.parent / "output" / "charts"
CHART_DIR.mkdir(parents=True, exist_ok=True)

PORTRAIT_WIDTH = 16.5  # cm
LANDSCAPE_WIDTH = 25.0  # cm

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_borders(cell, color="999999", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ("top","left","bottom","right"):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)

def set_table_fixed_layout(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

def section_landscape(doc):
    new = doc.add_section()
    new.orientation = WD_ORIENT.LANDSCAPE
    new.page_width, new.page_height = new.page_height, new.page_width
    new.left_margin = Cm(1.5); new.right_margin = Cm(1.5)
    new.top_margin = Cm(1.8); new.bottom_margin = Cm(1.8)

def section_portrait(doc):
    new = doc.add_section()
    new.orientation = WD_ORIENT.PORTRAIT
    if new.page_width > new.page_height:
        new.page_width, new.page_height = new.page_height, new.page_width
    new.left_margin = Cm(2); new.right_margin = Cm(2)
    new.top_margin = Cm(2); new.bottom_margin = Cm(2)

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor.from_string(color)
    return h

def add_para(doc, text, bold=False, size=10.5, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)
    return p

def add_callout(doc, title, body_lines, fill="FFF7E1", border="EFB54C"):
    """Yellow callout block for diagnosis + suggestions."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.width = Cm(PORTRAIT_WIDTH)
    set_cell_bg(cell, fill)
    set_cell_borders(cell, border, "8")
    p = cell.paragraphs[0]
    r = p.add_run(f"🔍 {title}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("8A5A00")
    for line in body_lines:
        para = cell.add_paragraph()
        run = para.add_run("• " + line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string("5A3A00")
    doc.add_paragraph()

def add_table(doc, headers, rows, widths_cm=None):
    """Render a table. headers: list[str]. rows: list[list[str]]. Auto-sizes if no widths."""
    if not rows:
        return None  # skip empty
    n = len(headers)
    table = doc.add_table(rows=1+len(rows), cols=n)
    table.style = "Light Grid Accent 1"
    set_table_fixed_layout(table)
    # Header
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ""
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string("FFFFFF")
        set_cell_bg(c, "1F3A93")
        if widths_cm and j < len(widths_cm):
            c.width = Cm(widths_cm[j])
    # Body
    for i, row in enumerate(rows):
        for j, val in enumerate(row[:n]):
            c = table.cell(i+1, j)
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            if widths_cm and j < len(widths_cm):
                c.width = Cm(widths_cm[j])
            # Color the dot column (last) if it's a status emoji
            if j == n - 1 and str(val).strip() in ("🟢", "🟡", "🔴"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r.font.size = Pt(14)
    doc.add_paragraph()
    return table

def chart_gmv_bar(this_week, last_week, this_month, last_month, currency, out_path):
    """Two-grouped bar chart: thisWeek/lastWeek + thisMonth/lastMonth."""
    if not HAS_PLT: return None
    labels = ["本周 / 上周", "本月 MTD / 上月全月"]
    this_vals = [this_week or 0, this_month or 0]
    prev_vals = [last_week or 0, last_month or 0]
    fig, ax = plt.subplots(figsize=(8.5, 4))
    x = range(len(labels))
    w = 0.36
    ax.bar([i-w/2 for i in x], this_vals, w, label="当前周期", color="#1F77B4")
    ax.bar([i+w/2 for i in x], prev_vals, w, label="对比周期", color="#FF7F0E")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels)
    ax.set_ylabel(f"GMV ({currency})")
    ax.set_title("GMV 周/月对比")
    ax.legend()
    for i, v in enumerate(this_vals):
        ax.text(i-w/2, v, f"{currency}{v:,.0f}", ha="center", va="bottom", fontsize=9)
    for i, v in enumerate(prev_vals):
        ax.text(i+w/2, v, f"{currency}{v:,.0f}", ha="center", va="bottom", fontsize=9)
    plt.tight_layout()
    plt.savefig(out_path, dpi=140)
    plt.close()
    return out_path

def chart_daily_trend(daily_rows, currency, out_path):
    """Line chart of this-week vs last-week daily amounts."""
    if not HAS_PLT or not daily_rows: return None
    labels = [r[0].split(",")[0][:3] for r in daily_rows]  # short weekday
    def parse(v):
        if v == "—": return 0
        return float(str(v).replace(currency,"").replace(",","").strip() or 0)
    this_vals = [parse(r[1]) for r in daily_rows]
    prev_vals = [parse(r[2]) for r in daily_rows]
    fig, ax = plt.subplots(figsize=(9, 3.5))
    ax.plot(labels, this_vals, marker="o", label="本周", color="#1F77B4", linewidth=2)
    ax.plot(labels, prev_vals, marker="s", label="上周", color="#FF7F0E", linewidth=2, linestyle="--")
    ax.set_ylabel(f"日 GMV ({currency})")
    ax.set_title("本周 vs 上周日度 GMV 走势")
    ax.legend()
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.savefig(out_path, dpi=140)
    plt.close()
    return out_path

def chart_channel_donut(channel_mix, region_code, out_path):
    if not HAS_PLT or not channel_mix: return None
    # truncate long publisher names for legend
    labels = [(r[0][:22] + "…") if len(r[0]) > 22 else r[0] for r in channel_mix]
    sizes = [float(r[1].rstrip("%")) for r in channel_mix]
    if sum(sizes) <= 0: return None
    colors = ["#1F77B4","#FF7F0E","#2CA02C","#D62728","#9467BD","#8C564B","#E377C2"]
    legend_labels = [f"{lbl} — {s:.1f}%" for lbl, s in zip(labels, sizes)]
    fig, ax = plt.subplots(figsize=(8, 4.2))
    wedges, _ = ax.pie(sizes, startangle=90,
                       colors=colors[:len(sizes)],
                       wedgeprops=dict(width=0.42, edgecolor="white", linewidth=2))
    ax.set_title(f"{region_code} — Publisher GMV 占比（Awin 后台数据）", fontsize=12, pad=10)
    ax.legend(wedges, legend_labels, loc="center left", bbox_to_anchor=(1.0, 0.5),
              fontsize=10, frameon=False)
    plt.tight_layout()
    plt.savefig(out_path, dpi=140, bbox_inches="tight")
    plt.close()
    return out_path

def build(data, out_path):
    doc = Document()
    # base section margins
    s = doc.sections[0]
    s.left_margin = Cm(2); s.right_margin = Cm(2)
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)

    region = data["region_name"]
    code = data["region_code"]
    cur = data["currency_symbol"]
    date_str = data["report_date"]

    # Title
    title = doc.add_paragraph()
    r = title.add_run(f"Awin Rockbros {region} 联盟营销周报")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string("1F3A93")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sr = sub.add_run(f"Merchant ID: {data['merchant_id']}  |  报告日期: {date_str}")
    sr.font.size = Pt(11); sr.font.color.rgb = RGBColor.from_string("666666")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ============ §1 执行摘要 ============
    add_heading(doc, "1. 执行摘要", level=1, color="1F3A93")
    if data.get("exec_kpi"):
        add_para(doc, "本期核心 KPI 快照（数据来源：Awin Performance Over Time + Affiliate Performance 后台）", size=9.5, color="666666")
        add_table(doc, ["指标", "当前", "对比", "基准 / 上周期", "状态"],
                  data["exec_kpi"], widths_cm=[3.5, 3, 2.5, 4.5, 1.5])
    if data.get("exec_takeaways"):
        add_para(doc, "三大判断:", bold=True, size=11)
        for t in data["exec_takeaways"]:
            add_para(doc, "  • " + t, size=10.5)
    if data.get("exec_diagnosis"):
        add_callout(doc, "诊断与建议", [data["exec_diagnosis"]] + data.get("exec_suggestions", []))

    # ============ §2 GMV 表现深度 ============
    add_heading(doc, "2. GMV 表现深度", level=1, color="1F3A93")

    # 2.1 周/月 GMV 对比表
    if data.get("gmv_summary"):
        add_heading(doc, "2.1 周期对比", level=2)
        add_table(doc, ["周期", "当前", "对比周期", "环比", "状态"],
                  data["gmv_summary"], widths_cm=[5, 3.5, 3.5, 2.5, 1.5])

    # 2.1b GMV bar chart (周/月对比)
    perf_summary = data.get("gmv_summary", [])
    if perf_summary and HAS_PLT:
        def gv(row, col):
            s = str(row[col]).replace(cur,"").replace(",","").strip()
            try: return float(s)
            except: return None
        tw = gv(perf_summary[0], 1) if len(perf_summary) > 0 else None
        lw = gv(perf_summary[0], 2) if len(perf_summary) > 0 else None
        tm = gv(perf_summary[1], 1) if len(perf_summary) > 1 else None
        lm = gv(perf_summary[1], 2) if len(perf_summary) > 1 else None
        path = chart_gmv_bar(tw, lw, tm, lm, cur, str(CHART_DIR / f"{code}_gmv_compare.png"))
        if path:
            doc.add_picture(path, width=Cm(15))
            doc.add_paragraph()

    # 2.2 日度 GMV 走势 (本周 vs 上周)
    if data.get("gmv_trend_week"):
        add_heading(doc, "2.2 日度 GMV 走势（本周 vs 上周）", level=2)
        if HAS_PLT:
            path = chart_daily_trend(data["gmv_trend_week"], cur, str(CHART_DIR / f"{code}_daily_trend.png"))
            if path:
                doc.add_picture(path, width=Cm(15))
                doc.add_paragraph()
        add_table(doc, ["日期", "本周 GMV", "上周 GMV", "环比", "备注"],
                  data["gmv_trend_week"], widths_cm=[4.5, 3.5, 3.5, 2.5, 2])

    # 2.3 后台截图
    if data.get("perf_thisweek_screenshot") and Path(data["perf_thisweek_screenshot"]).exists():
        add_heading(doc, "2.3 Awin 后台 — 本周 Performance Over Time 截图", level=2)
        doc.add_picture(data["perf_thisweek_screenshot"], width=Cm(16))
        doc.add_paragraph()

    if data.get("gmv_diagnosis"):
        add_callout(doc, "GMV 诊断与建议",
                    [data["gmv_diagnosis"]] + data.get("gmv_suggestions", []))

    # ============ §3 Publisher 表现 ============
    add_heading(doc, "3. Publisher 表现", level=1, color="1F3A93")

    # 3.1 Channel / Publisher GMV 占比 (donut + table)
    if data.get("channel_mix"):
        add_heading(doc, "3.1 Publisher GMV 占比（Awin 后台数据）", level=2)
        if HAS_PLT:
            path = chart_channel_donut(data["channel_mix"], code, str(CHART_DIR / f"{code}_donut.png"))
            if path:
                doc.add_picture(path, width=Cm(13))
                doc.add_paragraph()
        add_table(doc, ["Publisher", "数量占比", "周 GMV", "状态"],
                  data["channel_mix"], widths_cm=[6, 3, 4, 1.5])

    # Switch to landscape for wide top-publisher table
    section_landscape(doc)

    if data.get("top_publishers"):
        add_heading(doc, "3.2 Top Publishers (按 GMV 排序)", level=2)
        add_table(doc, ["Publisher", "ID", "Clicks", "CVR", "订单", "周 GMV", "佣金"],
                  data["top_publishers"],
                  widths_cm=[7, 2.5, 2.5, 2, 2, 3.5, 3])

    if data.get("pub_diagnosis"):
        add_callout(doc, "Publisher 诊断与建议",
                    [data["pub_diagnosis"]] + data.get("pub_suggestions", []),
                    fill="FFF7E1", border="EFB54C")

    # Back to portrait
    section_portrait(doc)

    # ============ §4 本周新增 Publisher ============
    if data.get("new_publishers"):
        section_landscape(doc)
        add_heading(doc, f"4. 本周新增 Publisher（{data.get('new_publishers_count',0)} 家）", level=1, color="1F3A93")
        add_table(doc, ["Publisher", "Website", "类型", "行业", "加入日期"],
                  data["new_publishers"],
                  widths_cm=[5, 5, 4.5, 4.5, 3])
        if data.get("recruit_diagnosis"):
            add_callout(doc, "招募诊断与建议",
                        [data["recruit_diagnosis"]] + data.get("recruit_suggestions", []))
        section_portrait(doc)

    # ============ §5 下周行动计划 ============
    section_landscape(doc)
    add_heading(doc, "5. 下周行动计划", level=1, color="1F3A93")
    if data.get("actions_growth"):
        add_heading(doc, "5.1 增长动作", level=2)
        for a in data["actions_growth"]:
            add_para(doc, "• " + a, size=10.5)
    if data.get("actions_recruit"):
        add_heading(doc, "5.2 招募动作", level=2)
        for a in data["actions_recruit"]:
            add_para(doc, "• " + a, size=10.5)
    if data.get("actions_optimize"):
        add_heading(doc, "5.3 优化动作", level=2)
        for a in data["actions_optimize"]:
            add_para(doc, "• " + a, size=10.5)
    section_portrait(doc)

    # ============ §6 KPI 仪表盘 ============
    if data.get("kpi_dashboard"):
        add_heading(doc, "6. KPI 仪表盘", level=1, color="1F3A93")
        add_table(doc, ["指标", "当前", "目标", "差距", "动作"],
                  data["kpi_dashboard"], widths_cm=[3.5, 3, 3, 2.5, 4])

    # Footer
    doc.add_paragraph()
    f = doc.add_paragraph()
    fr = f.add_run("数据来源：Awin 后台（performance-over-time + affiliate-performance + partnerships）。"
                    "自动化采集：CellDigital Affiliate Team。")
    fr.font.size = Pt(8); fr.font.color.rgb = RGBColor.from_string("999999")
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    print(f"Wrote {out_path}")

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--data", required=True)
    ap.add_argument("--out", required=True)
    args = ap.parse_args()
    data = json.load(open(args.data))
    build(data, args.out)

if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Generate ROCKBROS Simplified Weekly Report — DOCX + HTML + PDF.

Five sections in 专业咨询中文 (McKinsey-style consulting Chinese):
  1. 新增 Publisher 列表
  2. 过去一周与本月销量数据 + 后台截图
  3. Top performing publisher 列表 + GMV
  4. Publisher 合作请求邮件追溯（脱敏 — 不暴露邮箱）
  5. 品牌侧需要协助的事项

Hard rules (per global feedback memory):
  - NEVER expose publisher email addresses → mask as `xx****@domain` or DROP column
  - Wide tables → landscape section + explicit Cm() column widths
  - Each top-level section starts on a new page

Outputs three formats from a single data source:
  ~/Downloads/Rockbros-简化周报-YYYY-MM-DD.docx
  ~/Downloads/Rockbros-简化周报-YYYY-MM-DD.html
  ~/Downloads/Rockbros-简化周报-YYYY-MM-DD.pdf

Usage:
    python3 generate_simple_report.py --data output/rockbros_simple_data.json \
        --out-dir ~/Downloads --date 2026-04-16
"""
import argparse
import asyncio
import base64
import html as html_lib
import json
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --------------------------- Style constants ---------------------------
NAVY = RGBColor(0x0B, 0x2E, 0x4F)
ACCENT = RGBColor(0xE7, 0x4C, 0x3C)
GOLD = RGBColor(0xC9, 0x8A, 0x2B)
GREEN = RGBColor(0x2C, 0x8A, 0x4F)
GREY = RGBColor(0x55, 0x5F, 0x6D)
LIGHT_HEX = "F4F6F8"
HEADER_HEX = "0B2E4F"
DIAG_HEX = "FFF8E1"
ZH_FONT = "PingFang SC"
EN_FONT = "Calibri"

PORTRAIT_WIDTH_CM = 17.0   # A4 portrait usable width (after 2cm margins)
LANDSCAPE_WIDTH_CM = 25.7  # A4 landscape usable width


# --------------------------- Privacy helper ---------------------------
EMAIL_RE = re.compile(r"([a-zA-Z0-9_.+-]+)@([a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)")

def mask_email(s):
    """xx****@domain.com pattern."""
    if not isinstance(s, str):
        return s
    def _mask(m):
        local, domain = m.group(1), m.group(2)
        if len(local) <= 2:
            return f"{local[0]}****@{domain}"
        return f"{local[:2]}****@{domain}"
    return EMAIL_RE.sub(_mask, s)


# --------------------------- docx helpers ---------------------------
def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_run_zh(run, text, size=11, bold=False, italic=False, color=None):
    run.text = text
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), ZH_FONT)
    rfonts.set(qn("w:ascii"), EN_FONT)
    rfonts.set(qn("w:hAnsi"), EN_FONT)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, bold=False, size=11, color=None, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("")
    set_run_zh(r, text, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1, color=NAVY):
    sizes = {1: 17, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("")
    set_run_zh(r, text, size=sizes.get(level, 11), bold=True, color=color)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), HEADER_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def _set_cell_width(cell, cm):
    """Set width on tcW + cell.width (python-docx is finicky)."""
    cell.width = Cm(cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tcW = tc_pr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tc_pr.append(tcW)
    tcW.set(qn("w:w"), str(int(Cm(cm))))
    tcW.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths_cm, accent_first_col=False, header_size=10, body_size=9):
    """Build a table with EXPLICIT widths (cm). widths_cm length must == headers length."""
    assert len(widths_cm) == len(headers), "widths must match header count"
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    # Disable autofit at XML level
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is not None:
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tbl_pr.append(layout)

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("")
        set_run_zh(run, h, size=header_size, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(cell, HEADER_HEX)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_width(cell, widths_cm[i])

    # Data rows
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run("")
            is_first = (ci == 0 and accent_first_col)
            set_run_zh(run, str(val), size=body_size,
                       bold=is_first,
                       color=NAVY if is_first else None)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_width(cell, widths_cm[ci])
            if ri % 2 == 0:
                shade_cell(cell, LIGHT_HEX)
    return table


def add_diagnosis(doc, text, label="🔍 诊断与建议", width_cm=PORTRAIT_WIDTH_CM):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    _set_cell_width(cell, width_cm)
    shade_cell(cell, DIAG_HEX)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r = p1.add_run("")
    set_run_zh(r, label, size=11, bold=True, color=GOLD)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run("")
    set_run_zh(r, text, size=10, color=GREY)
    doc.add_paragraph()


def add_screenshot(doc, path, caption=None, width_in=6.0):
    if not path or not os.path.exists(path):
        add_para(doc, f"[截图缺失：{path}]", italic=True, color=GREY, size=9)
        return
    try:
        doc.add_picture(path, width=Inches(width_in))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run("")
            set_run_zh(r, caption, size=9, italic=True, color=GREY)
    except Exception as exc:
        add_para(doc, f"[截图嵌入失败：{exc}]", italic=True, color=GREY, size=9)


def add_banner(doc, brand, report_date, period, operator):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    _set_cell_width(cell, PORTRAIT_WIDTH_CM)
    shade_cell(cell, HEADER_HEX)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(2)
    r = p1.add_run("")
    set_run_zh(r, f"{brand} 联盟营销 · 简化版周报",
               size=20, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run("")
    set_run_zh(r, f"报告周期：{period}    |    出具日期：{report_date}",
               size=10, color=RGBColor(0xE5, 0xE7, 0xEB))
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(0)
    r = p3.add_run("")
    set_run_zh(r, operator, size=9, italic=True, color=RGBColor(0xCB, 0xD5, 0xE1))
    doc.add_paragraph()


def set_orientation(doc, orientation, width_cm=21.0, height_cm=29.7):
    """Switch the LAST section to portrait/landscape."""
    section = doc.sections[-1]
    if orientation == "landscape":
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
    else:
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


def new_section(doc, orientation="portrait"):
    """Insert a new section break + set orientation."""
    new = doc.add_section()
    if orientation == "landscape":
        new.orientation = WD_ORIENTATION.LANDSCAPE
        new.page_width = Cm(29.7)
        new.page_height = Cm(21.0)
    else:
        new.orientation = WD_ORIENTATION.PORTRAIT
        new.page_width = Cm(21.0)
        new.page_height = Cm(29.7)
    new.top_margin = Cm(2.0)
    new.bottom_margin = Cm(2.0)
    new.left_margin = Cm(2.0)
    new.right_margin = Cm(2.0)


# --------------------------- Privacy: scrub data once ---------------------------
def scrub_emails(data):
    """Drop email column from section4_emails; mask any other accidental emails."""
    out = json.loads(json.dumps(data))  # deep copy
    # section4_emails: original schema = [name, email, subject, type, last, status, requested]
    # New schema (no email): [name, subject, type, last, status, requested]
    new_rows = []
    for row in out.get("section4_emails", []):
        if len(row) == 7:
            name, _email, subject, ctype, last, status, requested = row
            new_rows.append([name, subject, ctype, last, status, requested])
        else:
            new_rows.append([mask_email(c) for c in row])
    out["section4_emails"] = new_rows
    # Catch any stray emails in other free-text fields
    for key in ("section4_todo", "section5_brand_actions", "section4_assets"):
        if key in out:
            out[key] = [[mask_email(c) for c in row] for row in out[key]]
    for key in ("section2_diagnosis", "section3_diagnosis", "section5_summary"):
        if key in out and isinstance(out[key], str):
            out[key] = mask_email(out[key])
    return out


# --------------------------- Sections (DOCX) ---------------------------
def section_1_new_publishers(doc, data):
    add_heading(doc, "一、本期新增 Publisher 列表", level=1)
    summary = data["section1_summary"]
    add_para(doc,
             f"过去 30 天内 ROCKBROS 双站点（US + EU）共新增 {summary['total_count']} 家 Publisher，"
             f"其中 US 站 {summary['us_count']} 家、EU 站 {summary['eu_count']} 家。"
             f"渠道结构以 Coupon / 优惠码（{summary['by_type'][0][2]}）为绝对主导，"
             f"Editorial 内容评测与 Cashback 返现各占约 16-17%，符合 ROCKBROS 户外用品品类典型流量结构。",
             size=10, color=GREY)

    add_heading(doc, "1.1 Publisher 类型分布", level=2)
    add_table(doc,
              ["Publisher 类型", "新增家数", "占比"],
              summary["by_type"],
              widths_cm=[9.0, 4.0, 4.0],
              accent_first_col=True)

    add_heading(doc, "1.2 30 家新增 Publisher 明细", level=2)
    add_table(doc,
              ["Publisher 名称", "区域", "网站 / 流量入口", "类型", "加入日期"],
              data["section1_new_publishers"],
              widths_cm=[5.0, 1.5, 4.5, 3.5, 2.5],
              accent_first_col=True)

    add_diagnosis(doc,
                  "新增结构良性 — Coupon 类占比 46.7% 符合 ROCKBROS 性价比户外用品定位；"
                  "建议 7 日内对 30 家全部推送『欢迎邮件 + 产品 feed + banner + 测试佣金』四件套激活包，"
                  "把首单率目标定在 ≥ 60%（行业均值 35-40%）。")


def section_2_sales(doc, data):
    add_heading(doc, "二、过去一周与本月销量数据", level=1)
    add_para(doc,
             "本节呈现 ROCKBROS US（Awin 58007）与 EU（Awin 122456）两站点的"
             "周度与月度（MTD）核心绩效指标，并附 Awin 后台原图截图作为佐证。",
             size=10, color=GREY)

    add_heading(doc, "2.1 周度绩效（Last 7 Days）", level=2)
    sales = data["section2_sales"]
    add_table(doc, sales["weekly"][0], sales["weekly"][1:],
              widths_cm=[5.0, 4.0, 4.0, 4.0],
              accent_first_col=True)

    add_heading(doc, "2.2 月度绩效（Month-to-Date）", level=2)
    add_table(doc, sales["monthly_mtd"][0], sales["monthly_mtd"][1:],
              widths_cm=[5.0, 4.0, 4.0, 4.0],
              accent_first_col=True)
    add_diagnosis(doc, data["section2_diagnosis"])

    add_heading(doc, "2.3 后台截图佐证", level=2)
    shots = data["screenshots"]
    add_screenshot(doc, shots["us_home"],
                   caption="图 2-1  US 站点 Awin 后台主屏（Advertiser 58007）")
    add_screenshot(doc, shots["eu_home"],
                   caption="图 2-2  EU 站点 Awin 后台主屏（Advertiser 122456）")


def section_3_top_publishers(doc, data):
    add_heading(doc, "三、Top Performing Publisher 排行", level=1)
    add_para(doc,
             "下表为本期跨双站点合并的 Top 14 Publisher 绩效排行，按 GMV 贡献降序排列；"
             "EU 站点头部高度集中（AudienceRun 三个 placement 占 EU GMV 70%）、"
             "US 站点呈现单一 YIELDKIT 主导格局，存在结构性单点依赖风险，"
             "下一阶段需通过 Coupon / Editorial / Cashback 三类流量补位。",
             size=10, color=GREY)

    add_heading(doc, "3.1 Top 14 Publisher 明细", level=2)
    add_table(doc,
              ["排名", "Publisher 名称", "区域", "GMV", "类型", "下一步动作"],
              data["section3_top_publishers"],
              widths_cm=[1.2, 4.5, 1.8, 2.0, 3.0, 4.5],
              accent_first_col=True)

    add_heading(doc, "3.2 后台明细截图", level=2)
    shots = data["screenshots"]
    add_screenshot(doc, shots["us_publishers"],
                   caption="图 3-1  US 站 Publisher Performance 全图（Awin 58007）")
    add_screenshot(doc, shots["eu_publishers"],
                   caption="图 3-2  EU 站 Publisher Performance 全图（Awin 122456）")
    add_diagnosis(doc, data["section3_diagnosis"])


def section_4_emails(doc, data):
    """LANDSCAPE — wider tables. Email column already removed by scrub_emails()."""
    new_section(doc, orientation="landscape")
    add_heading(doc, "四、Publisher 合作请求邮件追溯（过去 30 天）", level=1)
    add_para(doc,
             "本节基于联盟营销邮箱过去 30 天往来邮件梳理，共识别出 9 个与 ROCKBROS 直接相关的"
             "Publisher 合作 / Integration / Collaboration 请求，并据此拆解为 10 项可执行 TODO 与"
             "8 项配套素材文件需求。"
             "为保护合作伙伴隐私，本节已脱敏处理 — 不外露 Publisher 联系邮箱；如需联络方式请向"
             "联盟营销绩效办公室申请内部清单。",
             size=10, color=GREY)

    add_heading(doc, "4.1 合作请求邮件清单（脱敏版）", level=2)
    add_table(doc,
              ["Publisher", "邮件主题", "合作类型", "最后往来", "当前状态", "请求素材 / 资料"],
              data["section4_emails"],
              widths_cm=[4.5, 5.5, 3.5, 2.0, 4.0, 6.0],
              accent_first_col=True)

    add_heading(doc, "4.2 拆解后的 TODO 清单（按优先级）", level=2)
    add_table(doc,
              ["优先级", "任务", "Owner", "Deadline", "前置条件", "预期影响"],
              data["section4_todo"],
              widths_cm=[2.5, 6.0, 3.5, 2.5, 5.5, 5.5],
              accent_first_col=True)

    add_heading(doc, "4.3 配套素材 / 文件需求（PDF / PPT / 图库）", level=2)
    assets = data["section4_assets"]
    add_table(doc, assets[0], assets[1:],
              widths_cm=[6.5, 5.0, 8.5, 5.5],
              accent_first_col=True)

    add_diagnosis(doc,
                  "Publisher 请求高度集中于三类资产：(1) Awin 产品 feed（Geizhals/Coupons.de/Idealo 必需）、"
                  "(2) 高分辨率品牌素材包（HBR/Echowise/vatago.de/Golden Shopping Days 通用）、"
                  "(3) 季度 Coupon 计划（DACH 优惠码站急需）。"
                  "三类素材任一缺失都将直接卡住 €3,000+/月 EU GMV 增量。"
                  "建议：本周内先解锁产品 feed + 素材包两项，下周完成 Coupon 月历与 Bidding Policy 起草。",
                  width_cm=LANDSCAPE_WIDTH_CM)


def section_5_brand_actions(doc, data):
    new_section(doc, orientation="landscape")
    add_heading(doc, "五、品牌侧需要协助的事项", level=1)
    add_para(doc, data["section5_summary"], size=10, color=GREY)

    add_heading(doc, "5.1 品牌侧行动清单（按优先级与时间窗）", level=2)
    add_table(doc,
              ["优先级 / 时间窗", "行动事项", "责任方", "Deadline", "依据 / 影响"],
              data["section5_brand_actions"],
              widths_cm=[4.0, 7.5, 4.0, 2.5, 7.7],
              accent_first_col=True)

    add_diagnosis(doc,
                  "上述 10 项品牌侧行动按优先级分为四档：⛳ 必做（P0，本周内 3 项）、"
                  "🚀 高优（P1，两周内 3 项）、📅 计划（P2，本月内 3 项）、🛡️ 长期（P3，本季度 1 项）。"
                  "P0 三项构成『救火三角』 — 缺一项就会出现 Publisher 流失或行业负向口碑；"
                  "P1 三项构成『增长三角』 — 完成后预计可带来 €3,000-5,000/月 EU GMV 增量。"
                  "Affiliate Marketing 团队将于每周一晨会同步进度看板。",
                  label="🎯 战略判断",
                  width_cm=LANDSCAPE_WIDTH_CM)


def section_footer(doc):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("")
    set_run_zh(r, "— 报告完 —", size=10, italic=True, color=GREY)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("")
    set_run_zh(r,
               "Cell Digital Technology Inc.  |  联盟营销绩效办公室  |  数据来源：Awin Advertiser Console（联系邮箱已脱敏）",
               size=8, italic=True, color=GREY)


def build_docx(data, out_path):
    doc = Document()
    set_orientation(doc, "portrait")
    add_banner(doc, data["brand"], data["report_date"], data["report_period"], data["operator"])
    section_1_new_publishers(doc, data)
    doc.add_page_break()
    section_2_sales(doc, data)
    doc.add_page_break()
    section_3_top_publishers(doc, data)
    section_4_emails(doc, data)   # injects landscape section break
    section_5_brand_actions(doc, data)  # injects landscape section break
    section_footer(doc)
    doc.save(out_path)


# --------------------------- HTML builder ---------------------------
def _img_to_data_uri(path):
    if not path or not os.path.exists(path):
        return None
    ext = Path(path).suffix.lower().lstrip(".")
    mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
            "webp": "image/webp", "gif": "image/gif"}.get(ext, "image/png")
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("ascii")
    return f"data:{mime};base64,{b64}"


def _esc(s):
    return html_lib.escape(str(s))


def _html_table(headers, rows, accent_first_col=False, klass="data"):
    out = [f'<table class="{klass}">']
    out.append("<thead><tr>" + "".join(f"<th>{_esc(h)}</th>" for h in headers) + "</tr></thead>")
    out.append("<tbody>")
    for row in rows:
        cells = []
        for ci, val in enumerate(row):
            cls = ' class="first"' if (ci == 0 and accent_first_col) else ""
            cells.append(f"<td{cls}>{_esc(val)}</td>")
        out.append("<tr>" + "".join(cells) + "</tr>")
    out.append("</tbody></table>")
    return "\n".join(out)


def _html_diagnosis(text, label="🔍 诊断与建议"):
    return f'''<div class="callout">
  <div class="callout-title">{_esc(label)}</div>
  <div class="callout-body">{_esc(text)}</div>
</div>'''


def _html_screenshot(path, caption=None, embed=True):
    src = _img_to_data_uri(path) if embed else path
    if not src:
        return f'<p class="missing">[截图缺失：{_esc(path)}]</p>'
    cap = f'<figcaption>{_esc(caption)}</figcaption>' if caption else ""
    return f'<figure><img src="{src}" alt=""/>{cap}</figure>'


HTML_CSS = """
:root {
  --navy: #0B2E4F;
  --gold: #C98A2B;
  --grey: #555F6D;
  --light: #F4F6F8;
  --diag: #FFF8E1;
  --accent: #E74C3C;
  --green: #2C8A4F;
}
* { box-sizing: border-box; }
body {
  font-family: "PingFang SC", "Microsoft YaHei", -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
  color: #1f2937;
  max-width: 1180px;
  margin: 0 auto;
  padding: 32px 40px;
  line-height: 1.62;
  background: #fff;
}
.banner {
  background: var(--navy);
  color: #fff;
  padding: 36px 28px;
  border-radius: 8px;
  text-align: center;
  margin-bottom: 32px;
}
.banner h1 { font-size: 28px; margin: 0 0 8px; font-weight: 700; }
.banner .meta { color: #E5E7EB; font-size: 14px; margin: 0; }
.banner .org { color: #CBD5E1; font-size: 12px; font-style: italic; margin-top: 6px; }
h2 {
  color: var(--navy);
  border-bottom: 3px solid var(--navy);
  padding-bottom: 8px;
  margin-top: 36px;
  font-size: 22px;
}
h3 {
  color: var(--navy);
  margin-top: 24px;
  font-size: 16px;
  border-left: 4px solid var(--navy);
  padding-left: 10px;
}
p.lede { color: var(--grey); font-size: 14px; }
table.data {
  width: 100%;
  border-collapse: collapse;
  font-size: 13px;
  margin: 12px 0 18px;
  table-layout: auto;
}
table.data th {
  background: var(--navy);
  color: #fff;
  padding: 10px 8px;
  text-align: left;
  font-weight: 600;
  font-size: 13px;
}
table.data td {
  padding: 9px 8px;
  border-bottom: 1px solid #e5e7eb;
  vertical-align: top;
}
table.data tbody tr:nth-child(even) td { background: var(--light); }
table.data td.first { font-weight: 600; color: var(--navy); }
.callout {
  background: var(--diag);
  border-left: 5px solid var(--gold);
  padding: 14px 18px;
  margin: 18px 0;
  border-radius: 4px;
}
.callout-title { font-weight: 700; color: var(--gold); font-size: 14px; margin-bottom: 6px; }
.callout-body { color: var(--grey); font-size: 13px; line-height: 1.6; }
figure {
  text-align: center;
  margin: 18px 0;
  page-break-inside: avoid;
}
figure img {
  max-width: 100%;
  border: 1px solid #e5e7eb;
  border-radius: 4px;
  box-shadow: 0 2px 8px rgba(0,0,0,.06);
}
figcaption { color: var(--grey); font-size: 12px; font-style: italic; margin-top: 6px; }
.missing { color: var(--grey); font-style: italic; font-size: 12px; }
.section { page-break-before: always; padding-top: 16px; }
.section:first-of-type { page-break-before: auto; }
footer {
  margin-top: 48px;
  padding-top: 16px;
  border-top: 1px solid #e5e7eb;
  text-align: center;
  color: var(--grey);
  font-size: 11px;
  font-style: italic;
}
@media print {
  body { padding: 0; max-width: none; }
  .section { page-break-before: always; }
}
"""


def build_html(data, out_path, embed_images=True):
    s = data
    summary = s["section1_summary"]
    sales = s["section2_sales"]
    shots = s["screenshots"]

    parts = []
    parts.append(f"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{_esc(s['brand'])} 联盟营销 · 简化版周报 · {_esc(s['report_date'])}</title>
<style>{HTML_CSS}</style>
</head>
<body>

<div class="banner">
  <h1>{_esc(s['brand'])} 联盟营销 · 简化版周报</h1>
  <p class="meta">报告周期：{_esc(s['report_period'])}    |    出具日期：{_esc(s['report_date'])}</p>
  <p class="org">{_esc(s['operator'])}</p>
</div>
""")

    # Section 1
    parts.append('<section class="section">')
    parts.append("<h2>一、本期新增 Publisher 列表</h2>")
    parts.append(f'<p class="lede">过去 30 天内 ROCKBROS 双站点（US + EU）共新增 {summary["total_count"]} 家 Publisher，'
                 f'其中 US 站 {summary["us_count"]} 家、EU 站 {summary["eu_count"]} 家。'
                 f'渠道结构以 Coupon / 优惠码（{summary["by_type"][0][2]}）为绝对主导，'
                 f'Editorial 内容评测与 Cashback 返现各占约 16-17%，符合 ROCKBROS 户外用品品类典型流量结构。</p>')
    parts.append("<h3>1.1 Publisher 类型分布</h3>")
    parts.append(_html_table(["Publisher 类型", "新增家数", "占比"], summary["by_type"], accent_first_col=True))
    parts.append("<h3>1.2 30 家新增 Publisher 明细</h3>")
    parts.append(_html_table(
        ["Publisher 名称", "区域", "网站 / 流量入口", "类型", "加入日期"],
        s["section1_new_publishers"], accent_first_col=True))
    parts.append(_html_diagnosis(
        "新增结构良性 — Coupon 类占比 46.7% 符合 ROCKBROS 性价比户外用品定位；"
        "建议 7 日内对 30 家全部推送『欢迎邮件 + 产品 feed + banner + 测试佣金』四件套激活包，"
        "把首单率目标定在 ≥ 60%（行业均值 35-40%）。"))
    parts.append("</section>")

    # Section 2
    parts.append('<section class="section">')
    parts.append("<h2>二、过去一周与本月销量数据</h2>")
    parts.append('<p class="lede">本节呈现 ROCKBROS US（Awin 58007）与 EU（Awin 122456）两站点的'
                 '周度与月度（MTD）核心绩效指标，并附 Awin 后台原图截图作为佐证。</p>')
    parts.append("<h3>2.1 周度绩效（Last 7 Days）</h3>")
    parts.append(_html_table(sales["weekly"][0], sales["weekly"][1:], accent_first_col=True))
    parts.append("<h3>2.2 月度绩效（Month-to-Date）</h3>")
    parts.append(_html_table(sales["monthly_mtd"][0], sales["monthly_mtd"][1:], accent_first_col=True))
    parts.append(_html_diagnosis(s["section2_diagnosis"]))
    parts.append("<h3>2.3 后台截图佐证</h3>")
    parts.append(_html_screenshot(shots["us_home"], "图 2-1  US 站点 Awin 后台主屏（Advertiser 58007）", embed_images))
    parts.append(_html_screenshot(shots["eu_home"], "图 2-2  EU 站点 Awin 后台主屏（Advertiser 122456）", embed_images))
    parts.append("</section>")

    # Section 3
    parts.append('<section class="section">')
    parts.append("<h2>三、Top Performing Publisher 排行</h2>")
    parts.append('<p class="lede">下表为本期跨双站点合并的 Top 14 Publisher 绩效排行，按 GMV 贡献降序排列；'
                 'EU 站点头部高度集中（AudienceRun 三个 placement 占 EU GMV 70%）、'
                 'US 站点呈现单一 YIELDKIT 主导格局，存在结构性单点依赖风险，'
                 '下一阶段需通过 Coupon / Editorial / Cashback 三类流量补位。</p>')
    parts.append("<h3>3.1 Top 14 Publisher 明细</h3>")
    parts.append(_html_table(
        ["排名", "Publisher 名称", "区域", "GMV", "类型", "下一步动作"],
        s["section3_top_publishers"], accent_first_col=True))
    parts.append("<h3>3.2 后台明细截图</h3>")
    parts.append(_html_screenshot(shots["us_publishers"], "图 3-1  US 站 Publisher Performance 全图", embed_images))
    parts.append(_html_screenshot(shots["eu_publishers"], "图 3-2  EU 站 Publisher Performance 全图", embed_images))
    parts.append(_html_diagnosis(s["section3_diagnosis"]))
    parts.append("</section>")

    # Section 4
    parts.append('<section class="section">')
    parts.append("<h2>四、Publisher 合作请求邮件追溯（过去 30 天）</h2>")
    parts.append('<p class="lede">本节基于联盟营销邮箱过去 30 天往来邮件梳理，共识别出 9 个与 ROCKBROS 直接相关的'
                 'Publisher 合作 / Integration / Collaboration 请求，并据此拆解为 10 项可执行 TODO 与'
                 '8 项配套素材文件需求。'
                 '为保护合作伙伴隐私，本节已脱敏处理 — 不外露 Publisher 联系邮箱；如需联络方式请向'
                 '联盟营销绩效办公室申请内部清单。</p>')
    parts.append("<h3>4.1 合作请求邮件清单（脱敏版）</h3>")
    parts.append(_html_table(
        ["Publisher", "邮件主题", "合作类型", "最后往来", "当前状态", "请求素材 / 资料"],
        s["section4_emails"], accent_first_col=True))
    parts.append("<h3>4.2 拆解后的 TODO 清单（按优先级）</h3>")
    parts.append(_html_table(
        ["优先级", "任务", "Owner", "Deadline", "前置条件", "预期影响"],
        s["section4_todo"], accent_first_col=True))
    parts.append("<h3>4.3 配套素材 / 文件需求（PDF / PPT / 图库）</h3>")
    assets = s["section4_assets"]
    parts.append(_html_table(assets[0], assets[1:], accent_first_col=True))
    parts.append(_html_diagnosis(
        "Publisher 请求高度集中于三类资产：(1) Awin 产品 feed（Geizhals/Coupons.de/Idealo 必需）、"
        "(2) 高分辨率品牌素材包（HBR/Echowise/vatago.de/Golden Shopping Days 通用）、"
        "(3) 季度 Coupon 计划（DACH 优惠码站急需）。"
        "三类素材任一缺失都将直接卡住 €3,000+/月 EU GMV 增量。"
        "建议：本周内先解锁产品 feed + 素材包两项，下周完成 Coupon 月历与 Bidding Policy 起草。"))
    parts.append("</section>")

    # Section 5
    parts.append('<section class="section">')
    parts.append("<h2>五、品牌侧需要协助的事项</h2>")
    parts.append(f'<p class="lede">{_esc(s["section5_summary"])}</p>')
    parts.append("<h3>5.1 品牌侧行动清单（按优先级与时间窗）</h3>")
    parts.append(_html_table(
        ["优先级 / 时间窗", "行动事项", "责任方", "Deadline", "依据 / 影响"],
        s["section5_brand_actions"], accent_first_col=True))
    parts.append(_html_diagnosis(
        "上述 10 项品牌侧行动按优先级分为四档：⛳ 必做（P0，本周内 3 项）、"
        "🚀 高优（P1，两周内 3 项）、📅 计划（P2，本月内 3 项）、🛡️ 长期（P3，本季度 1 项）。"
        "P0 三项构成『救火三角』 — 缺一项就会出现 Publisher 流失或行业负向口碑；"
        "P1 三项构成『增长三角』 — 完成后预计可带来 €3,000-5,000/月 EU GMV 增量。"
        "Affiliate Marketing 团队将于每周一晨会同步进度看板。",
        label="🎯 战略判断"))
    parts.append("</section>")

    parts.append("""<footer>
  Cell Digital Technology Inc.  |  联盟营销绩效办公室  |  数据来源：Awin Advertiser Console（联系邮箱已脱敏）
</footer>
</body></html>""")

    Path(out_path).write_text("\n".join(parts), encoding="utf-8")


# --------------------------- PDF (via Playwright Chromium) ---------------------------
async def html_to_pdf(html_path, pdf_path):
    from playwright.async_api import async_playwright
    file_url = f"file://{Path(html_path).resolve()}"
    async with async_playwright() as p:
        browser = await p.chromium.launch()
        page = await browser.new_page()
        await page.goto(file_url, wait_until="networkidle")
        await page.pdf(
            path=str(pdf_path),
            format="A4",
            margin={"top": "18mm", "bottom": "18mm", "left": "16mm", "right": "16mm"},
            print_background=True,
            prefer_css_page_size=False,
        )
        await browser.close()


# --------------------------- Markdown (Obsidian) ---------------------------
def _md_table(headers, rows):
    out = ["| " + " | ".join(str(h) for h in headers) + " |"]
    out.append("|" + "|".join(["---"] * len(headers)) + "|")
    for row in rows:
        out.append("| " + " | ".join(str(c).replace("|", "\\|").replace("\n", " ") for c in row) + " |")
    return "\n".join(out)


def build_markdown(data, out_path, screenshot_rel_dir="attachments"):
    s = data
    summary = s["section1_summary"]
    sales = s["section2_sales"]
    parts = []
    parts.append(f"# {s['brand']} 联盟营销 · 简化版周报")
    parts.append("")
    parts.append(f"> **报告周期：** {s['report_period']}  ·  **出具日期：** {s['report_date']}")
    parts.append(f"> **出具方：** {s['operator']}")
    parts.append("")
    parts.append("---")
    parts.append("")
    parts.append("## 一、本期新增 Publisher 列表")
    parts.append("")
    parts.append(f"过去 30 天 ROCKBROS 双站点共新增 **{summary['total_count']} 家** Publisher（US {summary['us_count']} + EU {summary['eu_count']}）。")
    parts.append("")
    parts.append("### 1.1 Publisher 类型分布")
    parts.append(_md_table(["Publisher 类型", "新增家数", "占比"], summary["by_type"]))
    parts.append("")
    parts.append("### 1.2 30 家新增 Publisher 明细")
    parts.append(_md_table(
        ["Publisher 名称", "区域", "网站 / 流量入口", "类型", "加入日期"],
        s["section1_new_publishers"]))
    parts.append("")
    parts.append("> 🔍 **诊断与建议** — 新增结构良性，Coupon 占比 46.7% 符合 ROCKBROS 性价比户外定位。建议 7 日内对 30 家全部推送『欢迎邮件 + 产品 feed + banner + 测试佣金』四件套激活包，首单率目标 ≥ 60%。")
    parts.append("")
    parts.append("## 二、过去一周与本月销量数据")
    parts.append("")
    parts.append("### 2.1 周度绩效（Last 7 Days）")
    parts.append(_md_table(sales["weekly"][0], sales["weekly"][1:]))
    parts.append("")
    parts.append("### 2.2 月度绩效（Month-to-Date）")
    parts.append(_md_table(sales["monthly_mtd"][0], sales["monthly_mtd"][1:]))
    parts.append("")
    parts.append(f"> 🔍 **诊断与建议** — {s['section2_diagnosis']}")
    parts.append("")
    parts.append("### 2.3 后台截图")
    parts.append(f"![[{screenshot_rel_dir}/us_home.png]]")
    parts.append("*图 2-1  US 站点 Awin 后台主屏*")
    parts.append("")
    parts.append(f"![[{screenshot_rel_dir}/eu_home.png]]")
    parts.append("*图 2-2  EU 站点 Awin 后台主屏*")
    parts.append("")
    parts.append("## 三、Top Performing Publisher 排行")
    parts.append("")
    parts.append(_md_table(
        ["排名", "Publisher 名称", "区域", "GMV", "类型", "下一步动作"],
        s["section3_top_publishers"]))
    parts.append("")
    parts.append(f"![[{screenshot_rel_dir}/us_publishers_full.png]]")
    parts.append("*图 3-1  US 站 Publisher Performance 全图*")
    parts.append("")
    parts.append(f"![[{screenshot_rel_dir}/eu_publishers_full.png]]")
    parts.append("*图 3-2  EU 站 Publisher Performance 全图*")
    parts.append("")
    parts.append(f"> 🔍 **诊断与建议** — {s['section3_diagnosis']}")
    parts.append("")
    parts.append("## 四、Publisher 合作请求邮件追溯（过去 30 天，脱敏版）")
    parts.append("")
    parts.append("> ⚠️ 已脱敏 — 不暴露 Publisher 联系邮箱；如需联络方式请向联盟营销绩效办公室申请内部清单。")
    parts.append("")
    parts.append("### 4.1 合作请求邮件清单")
    parts.append(_md_table(
        ["Publisher", "邮件主题", "合作类型", "最后往来", "当前状态", "请求素材 / 资料"],
        s["section4_emails"]))
    parts.append("")
    parts.append("### 4.2 TODO 清单（按优先级）")
    parts.append(_md_table(
        ["优先级", "任务", "Owner", "Deadline", "前置条件", "预期影响"],
        s["section4_todo"]))
    parts.append("")
    parts.append("### 4.3 配套素材 / 文件需求")
    assets = s["section4_assets"]
    parts.append(_md_table(assets[0], assets[1:]))
    parts.append("")
    parts.append("> 🔍 **诊断与建议** — Publisher 请求集中于三类资产：(1) Awin 产品 feed、(2) 高分辨率品牌素材包、(3) 季度 Coupon 计划。三类任一缺失都将卡住 €3,000+/月 EU GMV 增量。本周先解锁 feed + 素材包，下周完成 Coupon 月历与 Bidding Policy。")
    parts.append("")
    parts.append("## 五、品牌侧需要协助的事项")
    parts.append("")
    parts.append(s["section5_summary"])
    parts.append("")
    parts.append("### 5.1 品牌侧行动清单（按优先级与时间窗）")
    parts.append(_md_table(
        ["优先级 / 时间窗", "行动事项", "责任方", "Deadline", "依据 / 影响"],
        s["section5_brand_actions"]))
    parts.append("")
    parts.append("> 🎯 **战略判断** — P0『救火三角』缺一项就会出现 Publisher 流失或负向口碑；P1『增长三角』完成可带来 €3,000-5,000/月 EU GMV 增量。每周一晨会同步看板。")
    parts.append("")
    parts.append("---")
    parts.append("*Cell Digital Technology Inc.  |  联盟营销绩效办公室  |  数据来源：Awin Advertiser Console（联系邮箱已脱敏）*")
    parts.append("")
    parts.append("#affiliate #rockbros #weekly-report")
    Path(out_path).write_text("\n".join(parts), encoding="utf-8")


# --------------------------- Main ---------------------------
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", required=True)
    parser.add_argument("--out-dir", required=True)
    parser.add_argument("--date", required=True, help="YYYY-MM-DD for filename")
    parser.add_argument("--obsidian-dir", default=None,
                        help="Optional: also write Markdown + attachments to Obsidian vault folder")
    parser.add_argument("--skip-pdf", action="store_true")
    args = parser.parse_args()

    data_path = Path(args.data).expanduser()
    out_dir = Path(args.out_dir).expanduser()
    out_dir.mkdir(parents=True, exist_ok=True)

    with data_path.open("r", encoding="utf-8") as f:
        raw_data = json.load(f)

    # PRIVACY: scrub once, all renderers consume scrubbed copy
    data = scrub_emails(raw_data)

    base = f"Rockbros-简化周报-{args.date}"
    docx_path = out_dir / f"{base}.docx"
    html_path = out_dir / f"{base}.html"
    pdf_path = out_dir / f"{base}.pdf"

    build_docx(data, docx_path)
    print(f"OK  DOCX → {docx_path}  ({docx_path.stat().st_size / 1024:.1f} KB)")

    build_html(data, html_path, embed_images=True)
    print(f"OK  HTML → {html_path}  ({html_path.stat().st_size / 1024:.1f} KB)")

    if not args.skip_pdf:
        try:
            asyncio.run(html_to_pdf(html_path, pdf_path))
            print(f"OK  PDF  → {pdf_path}  ({pdf_path.stat().st_size / 1024:.1f} KB)")
        except Exception as exc:
            print(f"WARN  PDF render failed: {exc}", file=sys.stderr)

    if args.obsidian_dir:
        ob_dir = Path(args.obsidian_dir).expanduser()
        ob_dir.mkdir(parents=True, exist_ok=True)
        attach_dir = ob_dir / "attachments"
        attach_dir.mkdir(exist_ok=True)
        # Copy screenshots into attachments
        for shot_key, shot_path in data["screenshots"].items():
            sp = Path(shot_path)
            if sp.exists():
                dest = attach_dir / sp.name
                dest.write_bytes(sp.read_bytes())
        md_path = ob_dir / f"{base}.md"
        build_markdown(data, md_path, screenshot_rel_dir="attachments")
        print(f"OK  MD   → {md_path}  ({md_path.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    main()

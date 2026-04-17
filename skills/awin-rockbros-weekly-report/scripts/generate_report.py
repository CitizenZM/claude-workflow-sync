#!/usr/bin/env python3
"""Generate Affiliate Weekly Report .docx from data.json.

Usage:
    python3 generate_report.py /path/to/data.json /path/to/output.docx [--screenshots dir]
"""
import json
import os
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---- Styling helpers ----
NAVY = RGBColor(0x0B, 0x2E, 0x4F)
ACCENT = RGBColor(0xE7, 0x4C, 0x3C)
GREY = RGBColor(0x55, 0x5F, 0x6D)
LIGHT_GREY_HEX = "F4F6F8"
HEADER_HEX = "0B2E4F"


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=NAVY):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, size=11, color=None, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    if color is not None:
        r.font.color.rgb = color
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        shade_cell(cell, HEADER_HEX)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if ri % 2 == 0:
                shade_cell(cell, LIGHT_GREY_HEX)
    if col_widths:
        for ri in range(len(table.rows)):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = w
    return table


def add_screenshot(doc, path, caption=None, width_in=6.2):
    if not path or not os.path.exists(path):
        return
    try:
        doc.add_picture(path, width=Inches(width_in))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = GREY
    except Exception as exc:
        add_para(doc, f"[screenshot embed failed: {exc}]", italic=True, color=GREY)


# ---- Cover + sections ----
def build_cover(doc, data):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("ROCKBROS GLOBAL")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT
    r.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Affiliate Weekly Report")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"

    rng = doc.add_paragraph()
    rng.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rng.add_run(f"Reporting Week ending {data['report_date']}  •  US (Awin 58007) + EU (Awin 122456)")
    r.font.size = Pt(12)
    r.font.color.rgb = GREY
    r.font.name = "Calibri"

    doc.add_paragraph()
    add_para(doc, "Prepared by:", bold=True, size=10, color=GREY)
    add_para(doc, "Cell Digital Technology Inc. — Affiliate Performance Office", size=10, color=GREY)
    add_para(doc, f"Source systems: Awin advertiser dashboards (US 58007, EU 122456). Data captured {data['report_date']}.",
             size=10, color=GREY, italic=True)
    doc.add_page_break()


def section_executive_summary(doc, data):
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(doc, "CEO / CMO / Growth Lead read time ≤ 2 minutes. Snapshot of consolidated US + EU affiliate performance, week-over-week movement, and top three takeaways.",
             italic=True, color=GREY, size=10)

    headers = ["Metric", "This Week", "WoW %", "Target", "Status"]
    add_table(doc, headers, data["exec_kpi"])

    doc.add_paragraph()
    add_heading(doc, "Key Takeaways", level=2)
    for tk in data["exec_takeaways"]:
        add_bullet(doc, tk)
    doc.add_page_break()


def section_gmv(doc, data):
    add_heading(doc, "2. GMV & Performance Deep Dive", level=1)

    add_heading(doc, "2.1 GMV Trend (US + EU consolidated)", level=2)
    add_table(doc, ["Day", "US GMV", "EU GMV", "Total", "Notes"], data["gmv_trend"])

    if data.get("us_home_screenshot"):
        add_screenshot(doc, data["us_home_screenshot"], caption="Awin US (58007) — advertiser home dashboard")
    if data.get("eu_home_screenshot"):
        add_screenshot(doc, data["eu_home_screenshot"], caption="Awin EU (122456) — advertiser home dashboard")

    add_heading(doc, "2.2 Channel Breakdown", level=2)
    add_table(doc, ["Channel Type", "GMV", "% Mix", "WoW", "Notes"], data["channel_breakdown"])
    add_para(doc, "Focus insight: balance between performance channels (cashback / coupon) and brand-building channels (content / creator). Target content GMV mix 30–50%.",
             italic=True, color=GREY, size=10)

    add_heading(doc, "2.3 Geo / Device / Product Split", level=2)
    add_table(doc, ["Dimension", "Top Driver", "Insight"], data["geo_device_product"])
    doc.add_page_break()


def section_publishers(doc, data):
    add_heading(doc, "3. Publisher Performance Analysis", level=1)

    add_heading(doc, "3.1 Top Publishers", level=2)
    add_table(doc, ["Publisher", "Region", "GMV", "WoW", "Type", "Action"], data["top_publishers"])

    if data.get("us_pub_screenshot"):
        add_screenshot(doc, data["us_pub_screenshot"], caption="Awin US — publisher list")
    if data.get("eu_pub_screenshot"):
        add_screenshot(doc, data["eu_pub_screenshot"], caption="Awin EU — publisher list")

    add_heading(doc, "3.2 Concentration Risk", level=2)
    add_table(doc, ["Metric", "Value", "Benchmark", "Risk Level"], data["concentration"])

    add_heading(doc, "3.3 Publisher Lifecycle Funnel", level=2)
    add_table(doc, ["Stage", "Count", "Conversion Rate"], data["lifecycle"])
    add_para(doc, "Diagnostic — low activation signals onboarding friction; low conversion signals offer / content mismatch.",
             italic=True, color=GREY, size=10)
    doc.add_page_break()


def section_recruitment(doc, data):
    add_heading(doc, "4. Recruitment & Activation", level=1)

    add_heading(doc, "4.1 Weekly Recruitment", level=2)
    add_table(doc, ["Source", "New Publishers", "Quality Score", "Notes"], data["recruitment"])

    add_heading(doc, "4.2 Activation Tracking", level=2)
    add_table(doc, ["Publisher", "Status", "First Content", "GMV", "Next Action"], data["activation"])
    add_para(doc, "KPIs: Time to first sale (TTFS), first-content → first-conversion lag.",
             italic=True, color=GREY, size=10)
    doc.add_page_break()


def section_campaigns(doc, data):
    add_heading(doc, "5. Campaign & Promotion Performance", level=1)

    add_heading(doc, "5.1 Campaign Results", level=2)
    add_table(doc, ["Campaign", "Type", "GMV", "ROI", "Notes"], data["campaigns"])

    add_heading(doc, "5.2 Offer Performance", level=2)
    add_table(doc, ["Offer Type", "CVR", "AOV", "Insight"], data["offers"])
    doc.add_page_break()


def section_content(doc, data):
    add_heading(doc, "6. Content & Creator Performance", level=1)

    add_heading(doc, "6.1 Content GMV Breakdown", level=2)
    add_table(doc, ["Content Type", "GMV", "% Mix", "WoW"], data["content_mix"])

    add_heading(doc, "6.2 Top Content", level=2)
    add_table(doc, ["Creator", "Platform", "Views", "GMV", "Hook"], data["top_content"])
    doc.add_page_break()


def section_risks(doc, data):
    add_heading(doc, "7. Issues & Risks", level=1)
    add_table(doc, ["Issue", "Impact", "Root Cause", "Action"], data["risks"])
    doc.add_page_break()


def section_actions(doc, data):
    add_heading(doc, "8. Next Week Action Plan", level=1)
    add_para(doc, "The single most important section. Each action has an owner and a measurable outcome.",
             italic=True, color=GREY, size=10)

    add_heading(doc, "8.1 Growth Actions", level=2)
    for a in data["actions_growth"]:
        add_bullet(doc, a)

    add_heading(doc, "8.2 Recruitment Plan", level=2)
    for a in data["actions_recruit"]:
        add_bullet(doc, a)

    add_heading(doc, "8.3 Optimization", level=2)
    for a in data["actions_optimize"]:
        add_bullet(doc, a)
    doc.add_page_break()


def section_kpi(doc, data):
    add_heading(doc, "9. KPI Dashboard (Operator View)", level=1)
    add_table(doc, ["KPI", "Current", "Target", "Gap", "Action"], data["kpi_dashboard"])
    add_para(doc, "Reviewed weekly. Color-code: green = on/above target, amber = within 10%, red = >10% gap.",
             italic=True, color=GREY, size=10)


def section_appendix(doc, data):
    doc.add_page_break()
    add_heading(doc, "Appendix — Source Captures", level=1)
    for sc in data.get("appendix_screenshots", []):
        if os.path.exists(sc.get("path", "")):
            add_screenshot(doc, sc["path"], caption=sc.get("caption", ""))
    add_para(doc, f"Generated {data['report_date']} via awin-rockbros-weekly-report skill (Cell Digital Technology Inc.).",
             italic=True, color=GREY, size=9)


def build(data, out_path):
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    build_cover(doc, data)
    section_executive_summary(doc, data)
    section_gmv(doc, data)
    section_publishers(doc, data)
    section_recruitment(doc, data)
    section_campaigns(doc, data)
    section_content(doc, data)
    section_risks(doc, data)
    section_actions(doc, data)
    section_kpi(doc, data)
    section_appendix(doc, data)

    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 3:
        print("Usage: generate_report.py data.json out.docx", file=sys.stderr)
        sys.exit(1)
    data_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    data = json.loads(data_path.read_text())
    data.setdefault("report_date", date.today().isoformat())
    build(data, str(out_path))
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()

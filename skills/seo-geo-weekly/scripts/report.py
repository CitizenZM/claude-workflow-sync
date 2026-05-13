"""Generate weekly SEO+GEO report as .docx."""
import datetime as dt
import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt, RGBColor


def _h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def _kv_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        t.cell(i, 0).text = str(k)
        t.cell(i, 1).text = str(v)
    return t


def _set_landscape(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def render(state_dir: Path, out_path: Path, store_name: str):
    catalog = json.loads((state_dir / "catalog.json").read_text())
    audit = json.loads((state_dir / "audit.json").read_text())
    geo = (
        json.loads((state_dir / "geo.json").read_text())
        if (state_dir / "geo.json").exists() else None
    )
    lh = (
        json.loads((state_dir / "lighthouse.json").read_text())
        if (state_dir / "lighthouse.json").exists() else None
    )
    serp = (
        json.loads((state_dir / "serp.json").read_text())
        if (state_dir / "serp.json").exists() else None
    )
    fixes = (
        json.loads((state_dir / "proposed_fixes.json").read_text())
        if (state_dir / "proposed_fixes.json").exists() else None
    )
    competitors = (
        json.loads((state_dir / "competitors.json").read_text())
        if (state_dir / "competitors.json").exists() else None
    )

    doc = Document()
    _set_landscape(doc)

    today = dt.date.today()
    week_no = today.isocalendar()[1]

    # Title
    title = doc.add_heading(f"Shopify {store_name} — SEO + GEO Weekly Report", level=0)
    doc.add_paragraph(f"Week {week_no} · {today.strftime('%B %d, %Y')}").italic = True
    doc.add_paragraph(f"Store: {catalog['shop']['shop']['name']} · "
                      f"{catalog['shop']['shop']['primaryDomain']['url']} · "
                      f"{catalog['shop']['shop']['plan']['displayName']} plan")

    # ── Executive Summary
    _h(doc, "Executive Summary", 1)
    issues = audit["issues"]
    score_perf = (lh or [{}])[0].get("scores", {}).get("performance") if lh else None
    score_seo = (lh or [{}])[0].get("scores", {}).get("seo") if lh else None
    summary_rows = [
        ("Products in catalog", len(catalog["products"])),
        ("Collections", len(catalog["collections"])),
        ("Pages crawled (storefront)", audit["pages_crawled"]),
        ("Lighthouse Performance (homepage)", f"{score_perf}/100" if score_perf is not None else "n/a"),
        ("Lighthouse SEO (homepage)", f"{score_seo}/100" if score_seo is not None else "n/a"),
        ("Pages with title >60 chars (will truncate)", len(issues["title_too_long"])),
        ("Pages with multiple <h1>", len(issues["multiple_h1"])),
        ("Pages missing meta description", len(issues["missing_meta_desc"])),
        ("Broken pages", len(issues["broken"])),
        ("robots.txt present", audit["robots_present"]),
        ("llms.txt present", audit["llms_txt_present"]),
    ]
    if geo:
        for engine, s in geo["summary"].items():
            if s.get("n"):
                summary_rows.append((
                    f"GEO — {engine} brand citation rate",
                    f"{s.get('brand_cited_pct', 0)}% across {s['n']} prompts",
                ))
    if serp:
        rk = serp.get("rankings", [])
        ranked = [r for r in rk if r.get("target_position")]
        summary_rows.append(("SERP keywords ranked (top 30)", f"{len(ranked)} of {len(rk)} checked"))
    if competitors:
        peer = next(((d, i) for d, i in competitors.items() if i.get("sitemap_urls")), None)
        if peer:
            summary_rows.append((f"Sitemap URL gap vs {peer[0]}", f"{audit['pages_crawled']} vs {peer[1]['sitemap_urls']}"))
    _kv_table(doc, summary_rows)

    # ── Technical SEO
    _h(doc, "1. Technical & On-Page SEO", 1)
    doc.add_paragraph(
        f"Crawled {audit['pages_crawled']} URLs from {audit['sitemap_url']}. "
        f"Findings below are change-detected weekly."
    )
    _h(doc, "Top issues by category", 2)
    for k, label in [
        ("title_too_long", "Title too long (>60 chars)"),
        ("multiple_h1", "Multiple <h1> tags"),
        ("missing_meta_desc", "Missing meta description"),
        ("missing_title", "Missing <title>"),
        ("missing_h1", "Missing <h1>"),
        ("broken", "Broken (4xx/5xx/no-response)"),
        ("alt_text_missing", "Images missing alt text"),
    ]:
        urls = issues.get(k, [])
        if not urls:
            continue
        _h(doc, f"{label} — {len(urls)} pages", 3)
        for u in urls[:10]:
            doc.add_paragraph(u, style="List Bullet")
        if len(urls) > 10:
            doc.add_paragraph(f"… and {len(urls) - 10} more", style="List Bullet")

    # ── JSON-LD Structured Data
    _h(doc, "Structured data (JSON-LD) coverage", 2)
    types = Counter()
    for p in audit.get("pages", []):
        for t in (p.get("json_ld_types") or []):
            if t:
                types[str(t)] += 1
    if types:
        t = doc.add_table(rows=1, cols=2)
        t.style = "Light Grid Accent 1"
        t.cell(0, 0).text = "Schema Type"
        t.cell(0, 1).text = "Pages"
        for k, n in types.most_common():
            row = t.add_row()
            row.cells[0].text = k
            row.cells[1].text = str(n)

    # ── Lighthouse / CWV
    if lh:
        _h(doc, "2. Core Web Vitals (Lighthouse)", 1)
        for r in lh:
            if r.get("error"):
                doc.add_paragraph(f"{r['url']} — error: {r['error']}")
                continue
            doc.add_paragraph(r["url"], style="Heading 3")
            scores = r.get("scores", {})
            vitals = r.get("vitals", {})
            _kv_table(doc, [
                ("Performance", f"{scores.get('performance')}/100"),
                ("SEO", f"{scores.get('seo')}/100"),
                ("Accessibility", f"{scores.get('accessibility')}/100"),
                ("Best Practices", f"{scores.get('best_practices')}/100"),
                ("LCP (target <2.5s)", vitals.get("lcp")),
                ("CLS (target <0.1)", vitals.get("cls")),
                ("TBT (target <200ms)", vitals.get("tbt")),
                ("FCP", vitals.get("fcp")),
                ("Speed Index", vitals.get("si")),
            ])

    # ── GEO
    if geo:
        _h(doc, "3. Generative Engine Optimization (GEO)", 1)
        doc.add_paragraph(
            "How the brand surfaces in AI answer engines. Lower citation rate = lower AI-driven discovery."
        )
        for engine, runs in geo["details"].items():
            _h(doc, engine.upper(), 2)
            sm = geo["summary"].get(engine, {})
            if sm.get("n"):
                _kv_table(doc, [
                    ("Prompts probed", sm["n"]),
                    ("Brand citation rate", f"{sm.get('brand_cited_pct', 0)}%"),
                    ("Competitor mentions (total)", sm.get("competitor_cited_total", 0)),
                    ("Avg brand position in list", sm.get("avg_brand_position") or "—"),
                ])
            for r in runs:
                if "error" in r:
                    continue
                p = doc.add_paragraph()
                p.add_run(f"Prompt: {r['prompt']}").bold = True
                cit = r["citations"]
                doc.add_paragraph(
                    f"  Brand cited: {'YES' if cit['brand_cited'] else 'NO'} · "
                    f"Position: {cit['brand_position_in_list'] if cit['brand_position_in_list'] > 0 else '—'} · "
                    f"Competitors mentioned: {', '.join(cit['competitors_cited']) or '—'}"
                )
                doc.add_paragraph(r["answer"][:1200] + ("…" if len(r["answer"]) > 1200 else ""))

    # ── Catalog hygiene
    _h(doc, "4. Catalog Hygiene (Admin API)", 1)
    miss_seo_t = [p for p in catalog["products"] if not (p.get("seo") or {}).get("title")]
    miss_seo_d = [p for p in catalog["products"] if not (p.get("seo") or {}).get("description")]
    miss_alt = [
        (p["title"], img["url"])
        for p in catalog["products"]
        for img in (p.get("images") or {}).get("nodes", [])
        if not img.get("altText")
    ]
    _kv_table(doc, [
        ("Total products", len(catalog["products"])),
        ("Products missing SEO title override", len(miss_seo_t)),
        ("Products missing SEO description override", len(miss_seo_d)),
        ("Product images missing alt text", len(miss_alt)),
    ])

    # ── SERP rankings
    if serp:
        _h(doc, "5. SERP Rankings (DuckDuckGo sample)", 1)
        rankings = serp.get("rankings", [])
        ranked = [r for r in rankings if r.get("target_position")]
        _kv_table(doc, [
            ("Keywords checked", len(rankings)),
            ("Brand ranked top 30", len(ranked)),
            ("Top 10", sum(1 for r in ranked if r["target_position"] <= 10)),
            ("Top 3", sum(1 for r in ranked if r["target_position"] <= 3)),
        ])
        if rankings:
            t = doc.add_table(rows=1, cols=5)
            t.style = "Light Grid Accent 1"
            for i, h in enumerate(["Keyword", "Our pos", "Δ WoW", "Top 10 domains", "Competitor positions"]):
                t.cell(0, i).text = h
            for r in rankings[:25]:
                row = t.add_row().cells
                row[0].text = r["query"]
                row[1].text = str(r.get("target_position") or "—")
                row[2].text = str(r.get("delta") or "—")
                row[3].text = ", ".join((r.get("top10_domains") or [])[:5])
                comps = r.get("competitor_positions") or {}
                row[4].text = "; ".join(f"{c}={p}" for c, p in comps.items() if p)

    # ── Competitor benchmarking
    if competitors:
        _h(doc, "6. Competitor Benchmarking", 1)
        t = doc.add_table(rows=1, cols=5)
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(["Domain", "Sitemap URLs", "Platform", "llms.txt", "Title pattern"]):
            t.cell(0, i).text = h
        for d, info in competitors.items():
            row = t.add_row().cells
            row[0].text = d
            row[1].text = str(info.get("sitemap_urls") or "—")
            row[2].text = ", ".join(info.get("platform") or []) or "—"
            row[3].text = "Yes" if info.get("llms_txt") else "No"
            row[4].text = (info.get("homepage_title") or "—")[:80]

    # ── Auto-fix proposals
    if fixes:
        _h(doc, "7. Auto-Fix Proposals", 1)
        doc.add_paragraph(
            f"Generated {len(fixes.get('rewritten_titles', []))} title rewrites, "
            f"{len(fixes.get('added_meta_descriptions', []))} meta descriptions, "
            f"{len(fixes.get('alt_text_proposals', []))} alt-text fills. "
            "Apply with `python3 run.py --apply` after review."
        )
        if fixes.get("rewritten_titles"):
            _h(doc, "Sample title rewrites", 2)
            t = doc.add_table(rows=1, cols=3)
            t.style = "Light Grid Accent 1"
            for i, h in enumerate(["Current", "Proposed", "Δ chars"]):
                t.cell(0, i).text = h
            for f in fixes["rewritten_titles"][:10]:
                if "error" in f: continue
                row = t.add_row().cells
                row[0].text = f.get("current_title", "")[:80]
                row[1].text = f.get("new_title", "")
                row[2].text = f"{f.get('current_len', 0)}→{f.get('new_len', 0)}"
        if fixes.get("theme_actions"):
            _h(doc, "Theme-level actions (manual)", 2)
            for a in fixes["theme_actions"]:
                doc.add_paragraph(f"⚠ {a['issue']}", style="Heading 3")
                doc.add_paragraph(f"Root cause: {a['root_cause']}")
                doc.add_paragraph(f"Fix: {a['fix']}")

    # ── Action items
    _h(doc, "8. Action Items (Prioritized)", 1)
    actions = []
    if len(issues["title_too_long"]) > 0:
        actions.append(("HIGH", f"Trim {len(issues['title_too_long'])} page titles to ≤60 chars to prevent SERP truncation."))
    if len(issues["multiple_h1"]) > 0:
        actions.append(("HIGH", f"Fix {len(issues['multiple_h1'])} pages with multiple <h1>; keep one primary heading per page."))
    if score_perf is not None and score_perf < 80:
        actions.append(("HIGH", f"Homepage Lighthouse Performance is {score_perf}/100. LCP and FCP both exceed targets — optimize hero image, defer non-critical JS, audit Shopify theme bloat."))
    if geo and any(s.get("brand_cited_pct", 0) == 0 for s in geo["summary"].values() if s.get("n")):
        actions.append(("HIGH", "Brand has 0% citation rate in AI engines. Build LLM-friendly content: detailed product FAQ pages, comparison guides, llms-full.txt with full catalog descriptions."))
    if len(miss_seo_t) > 0:
        actions.append(("MED", f"Set SEO title overrides for {len(miss_seo_t)} products."))
    if len(issues["missing_meta_desc"]) > 0:
        actions.append(("MED", f"Write meta descriptions for {len(issues['missing_meta_desc'])} pages."))
    if not audit["llms_txt_present"]:
        actions.append(("MED", "Publish /llms.txt and /llms-full.txt for AI crawler discoverability."))

    if actions:
        t = doc.add_table(rows=1, cols=2)
        t.style = "Light Grid Accent 1"
        t.cell(0, 0).text = "Priority"
        t.cell(0, 1).text = "Action"
        for prio, act in actions:
            r = t.add_row()
            r.cells[0].text = prio
            r.cells[1].text = act

    # ── Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(f"Generated {today.isoformat()} · seo-geo-weekly v1")
    r.italic = True
    r.font.size = Pt(8)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    import sys
    state_dir = Path(sys.argv[1])
    out = Path(sys.argv[2])
    store = sys.argv[3]
    print(render(state_dir, out, store))
